from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document

from data_model.enums import DocumentType


@dataclass
class ExtractedDocument:
    file_name: str
    document_type: DocumentType
    text: str
    page_count: int = 1


def extract_text(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".txt":
        return _extract_txt(path)
    raise ValueError("Unsupported file type. Supported: .pdf, .docx, .txt")


def _extract_pdf(path: Path) -> ExtractedDocument:
    doc = fitz.open(path)
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    text = "\n".join(pages).strip()
    doc.close()
    if not text:
        raise ValueError("PDF text extraction returned empty content.")
    return ExtractedDocument(
        file_name=path.name,
        document_type=DocumentType.PDF,
        text=text,
        page_count=len(pages),
    )


def _extract_docx(path: Path) -> ExtractedDocument:
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("DOCX text extraction returned empty content.")
    return ExtractedDocument(
        file_name=path.name,
        document_type=DocumentType.DOCX,
        text=text,
        page_count=1,
    )


def _extract_txt(path: Path) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        raise ValueError("TXT file is empty.")
    return ExtractedDocument(
        file_name=path.name,
        document_type=DocumentType.TXT,
        text=text,
        page_count=max(1, text.count("\f") + 1),
    )
