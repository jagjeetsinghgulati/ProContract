from __future__ import annotations

from pathlib import Path

from docx import Document

from data_model.schema import ContractModel


def generate_change_report(contract: ContractModel, output_path: str | Path) -> Path:
    document = Document()
    document.add_heading("ProContract Change Report", level=1)
    document.add_paragraph(f"Contract: {contract.file_name}")

    for clause in contract.clauses:
        if not clause.modifications:
            continue
        document.add_heading(clause.heading or clause.clause_id, level=2)
        document.add_paragraph("Original:")
        document.add_paragraph(clause.text)
        for proposal in clause.modifications:
            document.add_paragraph(f"Mode: {proposal.mode.value}")
            document.add_paragraph(f"Target: {proposal.target_party.value}")
            document.add_paragraph(f"Status: {proposal.accept_status}")
            document.add_paragraph("Proposed:")
            document.add_paragraph(proposal.proposed_text)
            document.add_paragraph(f"Summary: {proposal.change_summary}")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
