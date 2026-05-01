from __future__ import annotations

from pathlib import Path

from docx import Document

from data_model.schema import ContractModel


def export_modified_contract(contract: ContractModel, output_path: str | Path) -> Path:
    document = Document()
    document.add_heading("Modified Contract", level=1)
    document.add_paragraph(f"Source contract: {contract.file_name}")

    for clause in contract.clauses:
        document.add_heading(clause.heading or clause.clause_id, level=2)
        accepted = next(
            (m for m in clause.modifications if m.accept_status.lower() == "accepted"),
            None,
        )
        if accepted:
            document.add_paragraph("[MODIFIED]")
            document.add_paragraph(accepted.proposed_text)
        else:
            document.add_paragraph(clause.text)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
