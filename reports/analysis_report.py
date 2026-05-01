from __future__ import annotations

from pathlib import Path

from docx import Document

from data_model.schema import ContractModel


def generate_analysis_report(contract: ContractModel, output_path: str | Path) -> Path:
    document = Document()
    document.add_heading("ProContract Executive Analysis", level=1)
    document.add_paragraph(f"Contract: {contract.file_name}")
    document.add_paragraph(f"Clauses analyzed: {len(contract.clauses)}")
    document.add_paragraph("Parties:")
    for party in contract.parties:
        document.add_paragraph(f"- {party.display_name}", style="List Bullet")

    red_flags = [
        c for c in contract.clauses if c.classification and c.classification.red_flag
    ]
    document.add_heading("Key Risks", level=2)
    if not red_flags:
        document.add_paragraph("No red flags detected by the rule-based analyzer.")
    else:
        for clause in red_flags[:10]:
            document.add_paragraph(
                f"{clause.heading or clause.clause_id}: {clause.classification.rationale}",
                style="List Bullet",
            )

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
