from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document

from contract_parser.extract import extract_text
from data_model.enums import DocumentType


def test_extract_txt(tmp_path: Path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("This is a sample contract text.", encoding="utf-8")
    doc = extract_text(str(file_path))
    assert doc.document_type == DocumentType.TXT
    assert "sample contract" in doc.text


def test_extract_docx(tmp_path: Path):
    file_path = tmp_path / "sample.docx"
    docx = Document()
    docx.add_paragraph("Master Service Agreement")
    docx.add_paragraph("1. Payment Terms")
    docx.save(file_path)
    extracted = extract_text(str(file_path))
    assert extracted.document_type == DocumentType.DOCX
    assert "Master Service Agreement" in extracted.text


def test_extract_pdf(tmp_path: Path):
    file_path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Contract Title\n1. Termination Clause")
    pdf.save(file_path)
    pdf.close()

    extracted = extract_text(str(file_path))
    assert extracted.document_type == DocumentType.PDF
    assert "Termination Clause" in extracted.text
